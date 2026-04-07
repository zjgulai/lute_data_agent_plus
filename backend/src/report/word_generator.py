"""Word 报告生成器.

使用 python-docx 生成 .docx 格式的归因报告。
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT


class WordReportGenerator:
    """Word 报告生成器."""

    def generate(
        self,
        title: str,
        session_data: dict[str, Any],
        report_type: str,
    ) -> bytes:
        """生成 Word 报告并返回字节内容.

        Args:
            title: 报告标题
            session_data: 会话完整数据
            report_type: 报告类型 (process / full)

        Returns:
            docx 文件的字节内容
        """
        doc = Document()

        # 设置默认中文字体（尽量使用系统常见字体）
        from docx.oxml.ns import qn
        style = doc.styles["Normal"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        # 标题
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(
            f"会话ID: {session_data.get('session_id', '-')} | "
            f"分析周期: {session_data.get('start_date', '-')} ~ {session_data.get('end_date', '-')} | "
            f"生成时间: {session_data.get('generated_at', '-')}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x6B, 0x70, 0x80)

        # 一、分析概览
        doc.add_heading("一、分析概览", level=1)
        overview_table = doc.add_table(rows=1, cols=2)
        overview_table.style = "Light Grid Accent 1"
        hdr_cells = overview_table.rows[0].cells
        hdr_cells[0].text = "项目"
        hdr_cells[1].text = "内容"
        for key, value in [
            ("分析模式", session_data.get("analysis_mode", "-")),
            ("当前状态", session_data.get("current_state", "-")),
            ("对比周期", session_data.get("comparison_period", "-")),
        ]:
            row_cells = overview_table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)

        # 二、归因路径
        doc.add_heading("二、归因路径", level=1)
        steps = session_data.get("steps", [])
        if steps:
            path_table = doc.add_table(rows=1, cols=5)
            path_table.style = "Light Grid Accent 1"
            hdr = path_table.rows[0].cells
            hdr[0].text = "步骤"
            hdr[1].text = "节点"
            hdr[2].text = "类型"
            hdr[3].text = "选择维度"
            hdr[4].text = "熵减"

            for step in steps:
                row = path_table.add_row().cells
                row[0].text = f"Step {step.get('step_number', '-')}"
                row[1].text = step.get("node_name", "-")
                row[2].text = step.get("node_type", "-")
                selected = " / ".join(
                    filter(
                        None,
                        [step.get("selected_dimension") or "-", step.get("selected_child") or "-"],
                    )
                )
                row[3].text = selected
                entropy_results = step.get("entropy_results") or []
                if entropy_results:
                    er = entropy_results[0].get("entropy_reduction_normalized", 0)
                    row[4].text = f"{er * 100:.1f}%"
                else:
                    row[4].text = "-"
        else:
            doc.add_paragraph("暂无归因路径数据。")

        # 三、业务结论
        conclusion = session_data.get("conclusion")
        if conclusion:
            doc.add_heading("三、业务结论", level=1)
            conclusion_table = doc.add_table(rows=1, cols=2)
            conclusion_table.style = "Light Grid Accent 1"
            hdr_cells = conclusion_table.rows[0].cells
            hdr_cells[0].text = "项目"
            hdr_cells[1].text = "内容"

            confidence_map = {
                "high": "高",
                "medium": "中",
                "low": "低",
            }
            departments = conclusion.get("involved_departments") or []
            departments_str = ", ".join(departments) if isinstance(departments, list) else str(departments)

            for key, value in [
                ("原因类型", conclusion.get("reason_type", "-")),
                (
                    "置信度",
                    confidence_map.get(conclusion.get("confidence_level", ""), "-")
                    + f" ({conclusion.get('confidence_level', '-')})",
                ),
                ("涉及部门", departments_str or "-"),
                ("建议行动", conclusion.get("suggested_actions") or "-"),
                ("详细说明", conclusion.get("detailed_explanation", "-")),
            ]:
                row_cells = conclusion_table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)

        # 过程报告提示
        if report_type == "process":
            p = doc.add_paragraph()
            p_run = p.add_run(
                "说明：本报告为过程报告，分析尚未最终完成，结论部分可能为空。"
            )
            p_run.font.size = Pt(10)
            p_run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

        # 页脚
        doc.add_paragraph()
        footer = doc.add_paragraph("本报告由 GMV 智能归因系统自动生成，仅供内部决策参考。")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.runs[0]
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
